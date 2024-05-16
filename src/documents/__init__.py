from docx import Document as Docx
from docx.text.paragraph import Paragraph
from docx.shared import Pt
import re
import os
import json
import uuid
from itertools import chain

SECTIONS_TO_IGNORE = [
    'TITLE',
    'CICLO OTTO',
    'Conteúdos',
    'INTRODUÇÃO',
    'BIBLIOGRAFIA',
    'MECÂNICA DE VEÍCULOS LEVES',
    'SUMÁRIO',
    'INTRODUÇÃO',
    'METROLOGIA PARA MECÂNICA AUTOMOTIVA',
    'APRESENTAÇÃO',
    'REFERÊNCIAS',
    'REFERÊNCIAS BIBLIOGRÁFICAS',
    'Exercícios'
]

class Document():
    def __init__(self, file_path: str, sections_to_ignore: list[str] = []):
        self.document = Docx(file_path)
        self.file_path = file_path
        self.filename = os.path.basename(file_path)
        self.initial_heading = "TITLE"
        self.sections_to_ignore = sections_to_ignore
        self.__extract_contents()
        self.__set_id()

    def __set_id(self):
        if (not self.document.core_properties.identifier):
            self.document.core_properties.identifier = uuid.uuid4()
            self.document.save(self.file_path)

        self.id = self.document.core_properties.identifier

    def get_content(self) -> dict:
        return self.content

    def get_filename(self) -> str:
        return self.filename

    def __get_sections(self) -> list[str]:
        return list(self.content.keys())

    def get_id(self) -> int:
        return self.id

    def to_dict(self) -> list[dict]:
        output = []
        for section in self.__get_sections():
            output.append(dict(
                document_id=self.id,
                filename=self.get_filename(),
                section=section,
                content=" ".join(self.content[section]))
            )
        return output

    def __extract_contents(self):
        self.content = {}
        current_heading = self.initial_heading

        for section in self.document.sections:
            header_paragraphs = section.header.paragraphs
            footer_paragraphs = section.footer.paragraphs

            for item in section.iter_inner_content():
                if (not isinstance(item, Paragraph)):
                    continue

                text = self.__parse_text(item.text)
                if (len(text) == 0):
                    continue
                if (item in header_paragraphs or item in footer_paragraphs):
                    continue
                if (self.__has_small_font(item)):
                    continue

                if (self.__is_heading(item)):
                    current_heading = text

                section_content = self.content.setdefault(current_heading, [])

                if text != current_heading:
                    section_content.append(self.__parse_section_content(text))
        self.__remove_invalid_sections()

    def __parse_section_content(self, text: str) -> str:
        return text.replace("\"", "\\\"")

    def __remove_invalid_sections(self):
        if (isinstance(self.sections_to_ignore, list)):
            self.sections_to_ignore.append(self.initial_heading)

        for section in self.__get_sections():
            section_content_length = len(self.content[section])
            if section in self.sections_to_ignore or section_content_length == 0:
                self.content.pop(section)

    def __has_small_font(self, paragraph: Paragraph):
        font_size = paragraph.style.font.size

        if (self.__assert_font_size(font_size)):
            return True

        for run in paragraph.runs:
            if (self.__assert_font_size(run.font.size)):
                return True

        return False

    def __assert_font_size(self, font_size: Pt):
        return font_size == Pt(9) or font_size == Pt(8)

    def __is_heading(self, paragraph: Paragraph):
        return paragraph.style.name.startswith('Heading')

    def __parse_text(self, text: str):
        text = re.sub('\s+', ' ', text).strip()
        return '' if text.isnumeric() else text


class DocumentCollection():
    def __init__(self, dir_path: str, sections_to_ignore: list[str] = []) -> None:
        files = os.listdir(dir_path)
        self.documents = [Document(f'{dir_path}/{file_name}', sections_to_ignore)
                          for file_name in files]

    def get_documents(self) -> list[Document]:
        return self.documents

    def to_dict(self) -> list[dict]:
        return list(chain(*(document.to_dict() for document in self.documents)))

    def save(self, dest_dir: str) -> None:
        if not os.path.exists(dest_dir):
            os.makedirs(dest_dir)

        with open(f'{dest_dir}/raw_documents.json', 'w', encoding='utf-8') as fp:
            json.dump({document.get_filename(): document.get_content()
                      for document in self.documents}, fp, ensure_ascii=False)

        with open(f'{dest_dir}/documents.json', 'w', encoding='utf-8') as fp:
            json.dump(self.to_dict(), fp, ensure_ascii=False)
